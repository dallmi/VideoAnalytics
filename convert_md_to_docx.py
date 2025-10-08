#!/usr/bin/env python3
"""
Convert Markdown files to Word documents with proper native table support.
This script creates actual Word table objects instead of ASCII art.
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    """Set cell borders for better table formatting."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), kwargs[edge])
            edge_element.set(qn('w:sz'), '4')
            edge_element.set(qn('w:space'), '0')
            edge_element.set(qn('w:color'), '000000')
            tcBorders.append(edge_element)

    tcPr.append(tcBorders)


def parse_markdown_table(table_text):
    """Parse a markdown table into rows and columns."""
    lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]

    # Remove separator line (the one with dashes)
    lines = [line for line in lines if not re.match(r'^\s*\|?\s*[-:]+\s*(\|\s*[-:]+\s*)*\|?\s*$', line)]

    rows = []
    for line in lines:
        # Split by | and clean up
        cells = [cell.strip() for cell in line.split('|')]
        # Remove empty first/last cells from leading/trailing |
        if cells and not cells[0]:
            cells = cells[1:]
        if cells and not cells[-1]:
            cells = cells[:-1]
        if cells:
            rows.append(cells)

    return rows


def add_heading(doc, text, level=1):
    """Add a heading with proper formatting."""
    # Remove markdown heading symbols
    text = re.sub(r'^#{1,6}\s+', '', text)
    # Remove emojis and special symbols for cleaner headings
    text = re.sub(r'[✅❌⏸️▶️⏩⏪🔁📺😞🎛️🎮📹🎯📊🔑🎬📋📈🛠️🔍🧪📚🚀⚠️]', '', text)
    text = text.strip()

    heading = doc.add_heading(text, level=level)
    return heading


def convert_markdown_to_docx(md_file, docx_file):
    """Convert a markdown file to a Word document with native tables."""
    print(f"Converting {md_file} to {docx_file}...")

    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Add title from first # heading
    title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
    if title_match:
        doc.add_heading(title_match.group(1).strip(), 0)

    # Split content into sections
    lines = content.split('\n')
    i = 0

    in_code_block = False
    code_block_content = []
    in_table = False
    table_content = []

    while i < len(lines):
        line = lines[i]

        # Handle code blocks (preserve as monospace)
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                in_code_block = False
                if code_block_content:
                    p = doc.add_paragraph('\n'.join(code_block_content))
                    p.style = 'Normal'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                code_block_content = []
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # Detect markdown tables (lines with |)
        if '|' in line and not in_table:
            # Start collecting table
            in_table = True
            table_content = [line]
            i += 1
            continue

        if in_table:
            if '|' in line or re.match(r'^\s*[-:]+\s*$', line):
                table_content.append(line)
                i += 1
                continue
            else:
                # End of table, process it
                in_table = False
                table_rows = parse_markdown_table('\n'.join(table_content))

                if table_rows and len(table_rows) > 0:
                    # Create Word table
                    num_cols = len(table_rows[0])
                    table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    table.style = 'Light Grid Accent 1'

                    # Fill table
                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < num_cols:
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_text

                                # Format header row
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                                            run.font.size = Pt(10)
                                else:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.size = Pt(9)

                doc.add_paragraph()  # Add spacing after table
                table_content = []
                continue

        # Handle headings
        if line.startswith('#'):
            heading_level = len(re.match(r'^(#+)', line).group(1))
            if heading_level <= 3:  # Only include up to h3 in TOC
                add_heading(doc, line, min(heading_level, 9))
            i += 1
            continue

        # Handle regular paragraphs
        if line.strip():
            # Skip title (already added)
            if i == 0 and line.startswith('#'):
                i += 1
                continue

            # Add regular paragraph
            if not line.startswith('#') and not line.strip().startswith('```'):
                p = doc.add_paragraph(line)

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"✓ Created {docx_file}")


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python convert_md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)

    md_file = sys.argv[1]
    docx_file = sys.argv[2]

    convert_markdown_to_docx(md_file, docx_file)
