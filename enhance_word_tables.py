#!/usr/bin/env python3
"""
Post-process Word documents to enhance table formatting.
This script opens existing .docx files and improves table styling.
"""

import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def enhance_tables_in_docx(docx_path):
    """Enhance all tables in a Word document with better formatting."""
    print(f"Enhancing tables in {docx_path}...")

    try:
        doc = Document(docx_path)

        table_count = 0
        for table in doc.tables:
            table_count += 1

            # Apply a professional table style
            try:
                table.style = 'Table Grid'
            except:
                pass  # Skip if style not available

            # Format all cells
            for i, row in enumerate(table.rows):
                for cell in row.cells:
                    # Header row (first row)
                    if i == 0:
                        # Bold and slightly larger for headers
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(255, 255, 255)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        # Dark blue background for header
                        from docx.oxml.shared import OxmlElement, qn
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), '4472C4')  # Blue background
                        cell._element.get_or_add_tcPr().append(shading_elm)
                    else:
                        # Regular cells - smaller font
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        # Alternating row colors (light gray for even rows)
                        if i % 2 == 0:
                            from docx.oxml.shared import OxmlElement, qn
                            shading_elm = OxmlElement('w:shd')
                            shading_elm.set(qn('w:fill'), 'E7E6E6')  # Light gray
                            cell._element.get_or_add_tcPr().append(shading_elm)

            # Set column widths for better readability - NO WRAPPING in headers
            if len(table.columns) >= 4:
                # Make columns wider to prevent wrapping
                try:
                    from docx.shared import Inches
                    if len(table.columns) == 6:  # Tables with Action column
                        table.columns[0].width = Inches(1.4)  # timestamp - wider
                        table.columns[1].width = Inches(0.65)  # userId
                        table.columns[2].width = Inches(0.85)  # videoId
                        table.columns[3].width = Inches(1.1)  # eventName - wider
                        table.columns[4].width = Inches(0.9)  # currentTime - wider
                        table.columns[5].width = Inches(2.6)  # Action - much wider
                    elif len(table.columns) == 7:  # Tables with sessionId
                        table.columns[0].width = Inches(1.3)  # timestamp
                        table.columns[1].width = Inches(0.6)  # userId
                        table.columns[2].width = Inches(0.7)  # sessionId
                        table.columns[3].width = Inches(0.8)  # videoId
                        table.columns[4].width = Inches(1.0)  # eventName
                        table.columns[5].width = Inches(0.9)  # currentTime
                        table.columns[6].width = Inches(2.2)  # Action
                    elif len(table.columns) == 4:  # Simple 4-column tables
                        for col in table.columns:
                            col.width = Inches(1.5)
                except:
                    pass  # Skip if column width adjustment fails

        # Save the enhanced document
        doc.save(docx_path)
        print(f"✓ Enhanced {table_count} tables in {docx_path}")
        return True

    except Exception as e:
        print(f"✗ Error processing {docx_path}: {e}")
        return False


if __name__ == '__main__':
    if len(sys.argv) != 2:
        print("Usage: python enhance_word_tables.py <document.docx>")
        sys.exit(1)

    docx_file = sys.argv[1]
    success = enhance_tables_in_docx(docx_file)
    sys.exit(0 if success else 1)
